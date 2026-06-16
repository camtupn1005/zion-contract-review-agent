"""
ZION Contract Review Agent
--------------------------
Nhận 2 file hợp đồng (.docx hoặc .pdf), phân tích bằng Google Gemini,
xuất báo cáo rà soát tiếng Việt dạng .docx.

Cách dùng:
    python agent.py <file1> <file2> [--output <output_file>]

Ví dụ:
    python agent.py hop_dong_a.pdf hop_dong_b.docx
    python agent.py hop_dong_a.pdf hop_dong_b.docx --output BaoCao.docx
"""

import argparse
import os
import sys
from datetime import datetime
from pathlib import Path

import fitz  # PyMuPDF
import google.generativeai as genai
import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from dotenv import load_dotenv
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# Load env & config
# ─────────────────────────────────────────────

load_dotenv()

CONFIG_PATH = Path(__file__).parent / "config.yaml"

with open(CONFIG_PATH, "r", encoding="utf-8") as f:
    CONFIG = yaml.safe_load(f)


# ─────────────────────────────────────────────
# File reading utilities
# ─────────────────────────────────────────────

def read_docx(path: Path) -> str:
    """Trích xuất toàn bộ text từ file .docx."""
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def read_pdf(path: Path) -> str:
    """Trích xuất toàn bộ text từ file .pdf."""
    text_parts = []
    with fitz.open(path) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)


def extract_text(path: Path) -> str:
    """Đọc file hợp đồng, hỗ trợ .docx và .pdf."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return read_docx(path)
    elif suffix == ".pdf":
        return read_pdf(path)
    else:
        raise ValueError(f"Định dạng không được hỗ trợ: {suffix}. Chỉ hỗ trợ .docx và .pdf.")


# ─────────────────────────────────────────────
# Gemini API
# ─────────────────────────────────────────────

SYSTEM_PROMPT = """# AGENT: ZION – RÀ SOÁT HỢP ĐỒNG CTV
Bạn là hội đồng tư vấn pháp lý cao cấp của Công ty Cổ phần ZION (lĩnh vực trung gian thanh toán). Nhiệm vụ: rà soát hợp đồng giao khoán cộng tác viên, phát hiện rủi ro pháp lý, phân loại và routing đến đúng team nội bộ.

## VAI TRÒ – TƯ DUY BẮT BUỘC
Luôn mô phỏng đồng thời góc nhìn của:
Thanh tra lao động (Bộ LĐ-TB&XH) | Kiểm toán viên thuế (Tổng cục Thuế) | Điều tra viên BHXH | Thanh tra NHNN | Luật sư nguyên đơn tranh chấp lao động | Chuyên gia DD Big4 | Luật sư M&A | Kiểm toán viên nội bộ.

## NGUYÊN TẮC KHUNG (A1)

### 1. Chống QHLĐ trá hình (Điều 13, 16 BLLĐ VBHN 18/2026)
BẮT BUỘC loại bỏ: giờ làm việc/ca làm/lịch trực | chấm công/điểm danh/báo cáo định kỳ mang tính giám sát | mệnh lệnh/chỉ thị chi tiết | từ "lương/thưởng/phụ cấp" | "nhân viên/người lao động/biên chế" | quy trình nhân sự nội bộ.
YÊU CẦU: nhà thầu độc lập | thanh toán theo nghiệm thu kết quả | thời hạn HĐ ≤ 3 tháng | không giám sát quá trình | không làm tại địa điểm Bên A.

### 2. Nghĩa vụ thuế (Luật Thuế TNCN 2025 hiệu lực 1/7/2026 + TT 111/2013)
Bên giao khoán PHẢI khấu trừ TNCN tại nguồn (Điều 25 Luật Thuế TNCN 2025). Không thể từ bỏ nghĩa vụ này bằng HĐ dân sự. Khấu trừ 10% cho khoản TT ≥ 2 triệu/lần.

### 3. BHXH (Luật BHXH 2024)
Bên nhận khoán tự tham gia BHXH tự nguyện. HĐ giao khoán đúng pháp luật không phát sinh BHXH bắt buộc.

### 4. Công việc hợp pháp – KHÔNG giao khoán
Telesales tài chính | KYC/onboarding | thẩm định tín dụng | xử lý thanh toán/ủy thác | thu hồi nợ | rà soát gian lận.

## DANH SÁCH LUẬT ÁP DỤNG (A2)
1. BLLĐ (VBHN 18/2026) – Điều 13, 16, Chương III
2. BLDS 2015 – Điều 513–521, 562–569
3. Luật BHXH 2024
4. Luật Thuế TNCN 2025 (hiệu lực 1/7/2026) + TT 111/2013/TT-BTC
5. Luật Quản lý Thuế 2025 (hiệu lực 1/7/2026)
6. Luật Kế toán (VBHN 41/2026)
7. Luật SHTT (VBHN 67/2026)
8. Luật An ninh mạng 2025 (hiệu lực 1/7/2026)
9. Luật BVDLCN 2025 + NĐ 356/2025/NĐ-CP
10. Luật Đầu tư 2025 + NQ 66-17/2026/NQ-CP
11. TT 06/2023/TT-NHNN | NĐ 52/2024/NĐ-CP | NĐ 94/2025/NĐ-CP

## BẢNG PHÂN CÔNG TEAM PHỤ TRÁCH
Tech PO: Luồng kỹ thuật, sản phẩm, luồng TT/đối soát đặc thù; cam kết SLA
OP: Đối soát (thời hạn, quy trình); xử lý hoàn/hủy giao dịch
FPA: Chính sách phí dịch vụ; chi phí ZION trả đối tác; thỏa thuận tài chính đặc trưng
ACC: Thanh toán (thời hạn, quy trình, hồ sơ); tài khoản nhận TT; hóa đơn thuế
Risk: Nguồn tiền; gian lận/bồi hoàn; rủi ro kỹ thuật (appid, subaopid, binding); thẻ quốc tế
CPL: AML/CTF; đánh giá tuân thủ cho flow hợp tác mới
DP: Xử lý, chia sẻ dữ liệu cá nhân
LG: Hiệu lực/gia hạn/chấm dứt HĐ; vi phạm/bồi thường; bất khả kháng; tranh chấp; điều khoản chung
Biz: Điều khoản thương mại còn lại; quyết định khi đối tác không đồng ý điều chỉnh
Lưu ý: một điều khoản có thể thuộc nhiều team – liệt kê tất cả, team rủi ro cao hơn đứng trước.

## QUY TRÌNH THỰC HIỆN

### BƯỚC B – PRE-FLIGHT CHECK
B1. Kiểm tra tài liệu: cần đủ FILE CHUẨN (HĐ Mẫu gốc CTV-1) + FILE CẦN CHECK (HĐ thực tế). Xác nhận rõ vai trò từng file. Thiếu một trong hai → dừng, yêu cầu bổ sung.
B2. Kiểm tra hiệu lực luật A2. Phát hiện vấn đề → cảnh báo và chờ xác nhận. Pass toàn bộ → tiến hành Phần C ngay.

### BƯỚC C – RÀ SOÁT
Nhiệm vụ 1 – So sánh điều khoản KHÁC MẪU. Đánh giá theo 6 tiêu chí:
(1) Nội dung (2) Rủi ro công ty (3) Đa nghĩa (4) Tuân thủ pháp luật – trích dẫn điều luật cụ thể (5) Nguyên tắc khung A1 (6) Rủi ro tranh chấp.
Phân loại: 🔴 NGHIÊM TRỌNG | 🟡 TRUNG BÌNH | 🟢 THẤP
🔴 → [🚨 YÊU CẦU PHÊ DUYỆT CẤP TRÊN] | 🟡🟢 → Owner tự quyết
Nhiệm vụ 2 – Lỗi văn bản.
Nhiệm vụ 3 – Chỗ trống.
Nhiệm vụ 4 – Kiểm tra DLCN → yêu cầu Phụ lục CTV-2 + routing team DP.

### BƯỚC D – QUY TẮC BẮT BUỘC
- Đánh giá bản chất quan hệ thực tế, không chỉ ngôn ngữ HĐ
- Không bịa điều luật, không trích dẫn sai
- Kết quả chỉ gồm đúng 2 bảng tại Phần E
- Trả file .docx nếu > 3 vấn đề; ≤ 3 thì trả text

### BƯỚC E – KẾT QUẢ ĐẦU RA
KẾT QUẢ 1 – TÓM TẮT ĐIỀU HÀNH:
Mức độ rủi ro tổng thể | Khuyến nghị ký/không ký | Bảng: STT | Mức độ | Rủi ro | Phương án xử lý | Rủi ro nếu chấp nhận | Phân quyền | Team phụ trách

KẾT QUẢ 2 – BẢNG CHI TIẾT:
STT | Nhóm | Mức độ | Điều khoản | Nội dung HĐ Thực tế | Nội dung HĐ Mẫu | Đánh giá (6 tiêu chí) | Phương án xử lý | Team phụ trách
Sắp xếp: 🔴 → 🟡 → 🟢 → DLCN → LỖI VĂN BẢN → CHỖ TRỐNG

### BƯỚC F – CẬP NHẬT SAU XÁC NHẬN
Highlight VÀNG = chỗ trống | XANH LÁ = đã cập nhật.
Xuất file: [Tên HĐ]_REVIEWED_[YYYY-MM-DD]_v[n].docx.
Đính kèm Phụ lục BVDLCN nếu có DLCN."""


def build_prompt(contract1_text: str, contract2_text: str, sections: list) -> str:
    return f"""{SYSTEM_PROMPT}

---
FILE CHUẨN (HĐ Mẫu gốc CTV-1):
{contract1_text}

---
FILE CẦN CHECK (HĐ Thực tế):
{contract2_text}

---
Thực hiện đầy đủ Bước B → C → D → E theo quy trình trên. Xuất kết quả đúng 2 bảng tại Bước E.
"""


def call_gemini(prompt: str) -> str:
    api_key = os.getenv(CONFIG["llm"]["api_key_env"])
    if not api_key:
        raise EnvironmentError(
            f"Chưa thiết lập biến môi trường '{CONFIG['llm']['api_key_env']}'. "
            "Hãy tạo file .env và thêm GEMINI_API_KEY=<your_key>"
        )

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(
        model_name=CONFIG["llm"]["model"],
        generation_config={
            "temperature": CONFIG["llm"]["temperature"],
            "max_output_tokens": CONFIG["llm"]["max_output_tokens"],
        },
    )

    print(f"  → Đang gọi Gemini ({CONFIG['llm']['model']})...")
    response = model.generate_content(prompt)
    return response.text


# ─────────────────────────────────────────────
# Word report builder
# ─────────────────────────────────────────────

def add_horizontal_line(paragraph):
    """Thêm đường kẻ ngang dưới paragraph."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_docx_report(
    review_text: str,
    file1_name: str,
    file2_name: str,
    output_path: Path,
    sections: list,
):
    doc = Document()

    # ── Metadata / styles ──
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ── Cover ──
    title = doc.add_heading("BÁO CÁO RÀ SOÁT HỢP ĐỒNG", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(CONFIG["agent"]["name"])
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    add_horizontal_line(doc.add_paragraph())

    # ── Meta info ──
    info_para = doc.add_paragraph()
    info_para.add_run("Hợp đồng 1: ").bold = True
    info_para.add_run(file1_name)
    doc.add_paragraph().add_run("")

    info_para2 = doc.add_paragraph()
    info_para2.add_run("Hợp đồng 2: ").bold = True
    info_para2.add_run(file2_name)
    doc.add_paragraph().add_run("")

    info_para3 = doc.add_paragraph()
    info_para3.add_run("Ngày lập báo cáo: ").bold = True
    info_para3.add_run(datetime.now().strftime("%d/%m/%Y %H:%M"))

    add_horizontal_line(doc.add_paragraph())
    doc.add_page_break()

    # ── Review content ──
    # Tách nội dung theo từng section heading
    lines = review_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Detect heading (bắt đầu bằng số thứ tự hoặc ## / ###)
        is_main_heading = any(
            stripped.lower().startswith(s.lower()[:10])
            for s in sections
        ) or (len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in ".)")

        is_md_heading = stripped.startswith("##") or stripped.startswith("**")

        if is_main_heading or stripped.startswith("###"):
            h = doc.add_heading(stripped.lstrip("#").strip(), level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        elif stripped.startswith("##"):
            h = doc.add_heading(stripped.lstrip("#").strip(), level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        elif stripped.startswith("|"):
            # Bảng so sánh dạng markdown
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            # Kiểm tra separator row (---)
            if all(set(c) <= set("-: ") for c in cells):
                continue
            # Tìm table cuối, nếu chưa có thì tạo mới
            if not doc.tables or doc.tables[-1]._element.getnext() is not None:
                tbl = doc.add_table(rows=1, cols=len(cells))
                tbl.style = "Table Grid"
                row = tbl.rows[0]
                for i, cell_text in enumerate(cells):
                    row.cells[i].text = cell_text
                    for para in row.cells[i].paragraphs:
                        for run in para.runs:
                            run.bold = True
            else:
                tbl = doc.tables[-1]
                row = tbl.add_row()
                actual_cells = min(len(cells), len(row.cells))
                for i in range(actual_cells):
                    row.cells[i].text = cells[i]
        else:
            clean = stripped.lstrip("*").rstrip("*").strip()
            if stripped.startswith("**") and stripped.endswith("**"):
                p = doc.add_paragraph()
                p.add_run(clean).bold = True
            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(stripped[2:])
            else:
                doc.add_paragraph(stripped)

    doc.save(output_path)
    print(f"  ✓ Đã lưu báo cáo: {output_path}")


# ─────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="ZION Contract Review Agent — Rà soát và so sánh 2 hợp đồng"
    )
    parser.add_argument("file1", help="Đường dẫn hợp đồng thứ nhất (.docx hoặc .pdf)")
    parser.add_argument("file2", help="Đường dẫn hợp đồng thứ hai (.docx hoặc .pdf)")
    parser.add_argument(
        "--output", "-o",
        help="Tên file báo cáo đầu ra (.docx). Mặc định: BaoCao_RaSoatHopDong_<timestamp>.docx",
        default=None,
    )
    args = parser.parse_args()

    file1 = Path(args.file1)
    file2 = Path(args.file2)

    for f in [file1, file2]:
        if not f.exists():
            print(f"[LỖI] Không tìm thấy file: {f}")
            sys.exit(1)

    # Output path
    if args.output:
        output_path = Path(args.output)
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        prefix = CONFIG["output"]["filename_prefix"]
        output_path = Path(f"{prefix}_{timestamp}.docx")

    print("\n═══════════════════════════════════════════")
    print("  ZION Contract Review Agent")
    print("═══════════════════════════════════════════")
    print(f"  Hợp đồng 1 : {file1.name}")
    print(f"  Hợp đồng 2 : {file2.name}")
    print(f"  Đầu ra     : {output_path}")
    print("───────────────────────────────────────────")

    print("\n[1/3] Đọc nội dung hợp đồng...")
    text1 = extract_text(file1)
    text2 = extract_text(file2)
    print(f"  ✓ Hợp đồng 1: {len(text1):,} ký tự")
    print(f"  ✓ Hợp đồng 2: {len(text2):,} ký tự")

    print("\n[2/3] Phân tích bằng Gemini AI...")
    sections = CONFIG["review"]["sections"]
    prompt = build_prompt(text1, text2, sections)
    review_text = call_gemini(prompt)
    print(f"  ✓ Nhận được {len(review_text):,} ký tự phân tích")

    print("\n[3/3] Tạo báo cáo Word...")
    build_docx_report(review_text, file1.name, file2.name, output_path, sections)

    print("\n═══════════════════════════════════════════")
    print(f"  HOÀN THÀNH! Báo cáo: {output_path.resolve()}")
    print("═══════════════════════════════════════════\n")


if __name__ == "__main__":
    main()
